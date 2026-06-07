#!/usr/bin/env python3
"""Generate ATS-friendly DOCX and visually polished HTML/PDF resume."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from weasyprint import HTML

OUT_DIR = Path(__file__).resolve().parent

CONTACT = {
    "name": "Jason Jones",
    "title": "Systems Administrator & DevOps Engineer",
    "location": "Sioux Falls, SD",
    "phone": "303-885-2018",
    "email": "jasonjones21@att.net",
    "github": "https://github.com/RamenAnime",
    "portfolio": "https://ramenanime.com",
}

SUMMARY = (
    "Systems Administrator and DevOps Engineer with 12+ years of experience managing "
    "multi-site IT infrastructure, enterprise servers, high-availability networks, and "
    "cross-functional operations. Proven track record replacing manual workflows with "
    "custom automation, full-stack software, and CI/CD pipelines. Experienced in "
    "security incident response, SIEM deployment, firewall forensics, and third-party "
    "compliance coordination during active breach investigations. Strong bridge between "
    "engineering execution and frontline technical operations with a focus on reliability, "
    "uptime, and measurable operational outcomes."
)

ACHIEVEMENTS = [
    (
        "Led active security incident response at Kenworth of Jacksonville, coordinating "
        "with SonicWall vendor (Western NRG) and a third-party compliance team; analyzed "
        "firewall forensics across 9,000+ malware block events, identified attack initiator "
        "hosts, VPN bypass attempts, and data-harvesting browser extensions; directed Wazuh "
        "SIEM deployment across the environment."
    ),
    (
        "Built and launched a production-ready global marketplace platform with a 42-table "
        "relational database, real-time engines, and end-to-end cloud infrastructure as a "
        "solo technical project."
    ),
    (
        "Engineered a live digital dispatch board and internal ticketing system that "
        "replaced legacy paper processing, eliminating data lag and improving cross-department "
        "workflow tracking."
    ),
    (
        "Administered distributed multi-site IT systems across several locations with zero "
        "external vendor dependency, including hardware procurement, UniFi network "
        "architectures, and server provisioning."
    ),
    (
        "Coached and led a remote technical escalation team of 13 staff in a 24/7 "
        "mission-critical Tier 3 environment, handling 500+ advanced technical cases weekly "
        "while exceeding SLA and QA baselines."
    ),
]

EXPERIENCE = [
    {
        "title": "Systems Administrator, DevOps Engineer & Incident Responder",
        "company": "Kenworth of Jacksonville / PacLease",
        "location": "Garden City, GA",
        "dates": "2024 - Present",
        "bullets": [
            "Serve as sole IT Administrator overseeing multi-site production infrastructure across two regional locations; deployed UniFi network topologies, managed VLAN routing, and architected localized secure data access layers.",
            "Lead active cybersecurity incident response: analyzed SonicWall GMS firewall forensics identifying malware-contacting endpoints, VPN-bypass hosts, a data-harvesting browser extension (ntp2.mywavehome.net / wavebrowser), and an attack-initiating internal host; coordinate with firewall vendor and third-party compliance team on forensics timeline and scope. Direct Wazuh SIEM deployment for ongoing endpoint monitoring.",
            "Engineered, tested, and deployed a custom digital dispatch system, owning the full software lifecycle from development through production, eliminating business status lag and manual paper workflows.",
            "Perform complete system administration for the enterprise Dealer Management System (Karmak DMS), overseeing access management, data reporting, workflow compliance, and complex transaction tracking.",
            "Manage hardware lifecycle pipelines including bare-metal server configuration, workstation deployment, and client configuration, reducing technical response intervals from days to under two hours.",
            "Act as critical escalation point for infrastructure faults, vendor interconnections, and network-attached appliances, maintaining 100% operational availability across distributed logistics and parts facilities.",
        ],
    },
    {
        "title": "Systems Coordinator & Shop Foreman",
        "company": "MHC Kenworth",
        "location": "Savannah, GA",
        "dates": "2022 - 2024",
        "bullets": [
            "Promoted from service tracking to operations foreman to orchestrate technological execution, resource distribution, and technical workflows for a high-output maintenance floor supporting 120+ assets monthly.",
            "Managed internal resource planning matrices, aligned technician specialization with technical task profiles, and verified end-to-end quality and corporate system documentation compliance.",
            "Maintained enterprise platform inputs for all in-progress service pipelines, establishing high database entry accuracy and managing client authorizations for system updates.",
            "Designed technician instructional frameworks and technical standardization documentation that lowered comeback margins and aligned with warranty system requirements.",
        ],
    },
    {
        "title": "Site Operations Manager",
        "company": "Genuine Parts Company (IBS) / Liberty County School District",
        "location": "Hinesville, GA",
        "dates": "2019 - 2022",
        "bullets": [
            "Held end-to-end contractual P&L and technical inventory tracking liability for a Fortune 500-backed on-site operations initiative, administering high-volume procurement budgets and vendor invoicing.",
            "Supervised a technical support and inventory team of six, executing performance evaluations, professional skills tracking, and operational standard alignment across a multi-year tenure.",
            "Overhauled local asset tracking protocols to eliminate stock shortfalls, maintaining 90%+ same-day target achievement while consolidating multi-vendor pipelines to limit infrastructure overhead.",
        ],
    },
    {
        "title": "Service Operations Manager",
        "company": "Express Oil Change & Tire Engineers",
        "location": "Savannah, GA",
        "dates": "2017 - 2019",
        "bullets": [
            "Directed multi-shift technical facility operations, managing team schedules, workflow execution paths, systemic safety parameters, and customer-facing technical support channels.",
            "Elevated site throughput and customer validation trends through workflow standardization and direct SLA compliance.",
        ],
    },
    {
        "title": "Remote Technical Support Coach & Tier 3 Case Manager",
        "company": "DISH Network",
        "location": "Remote",
        "dates": "2012 - 2017",
        "bullets": [
            "Coached and directed a fully remote 13-person Tier 3 advanced technical support team in a 24/7 model handling 500+ complex escalated incidents per week.",
            "Monitored strict adherence to SLA baselines, qualitative performance metrics, and compliance scorecards; executed targeted technical coaching to lower recurring incident rates.",
            "Created comprehensive curriculum and conducted deep technical onboarding covering diagnostic procedures, incident recording standards, and compliance frameworks.",
        ],
    },
]

PROJECT = {
    "name": "Ramen Anime | Global Collectibles Marketplace & Social Platform",
    "links": "https://github.com/RamenAnime/RamenAnime-Portfolio | https://ramenanime.com",
    "bullets": [
        "Full-Stack & Architecture: Designed a production-grade web ecosystem featuring real-time auction bidding, multi-currency engines, and social elements on containerized cloud infrastructure.",
        "Database Design: Structured a 42-table relational database schema using Drizzle ORM on TiDB Cloud (MySQL-compatible).",
        "DevOps & Infrastructure: Git version control, Render automated hosting, GoDaddy DNS routing, Vitest/ESLint/Prettier CI/CD pipelines.",
        "Core Stack: React 19, TypeScript, Vite, Tailwind CSS, shadcn/ui, Node.js 24, tRPC 11, Hono, JWT, reCAPTCHA.",
    ],
}

SKILLS_ATS = [
    "Systems Administration: Windows Server, Linux (RHEL, Ubuntu, Arch), UniFi Networking, VLAN Routing, Bare-Metal & Cloud Provisioning, TrueNAS SCALE, CRM & Ticket Platforms",
    "DevOps & Automation: CI/CD Workflows, Git, GitHub, Containerized Deployments, Automation Engineering, Render, Python Scripting",
    "Security & Incident Response: SonicWall GMS, Firewall Forensics, Wazuh SIEM, Endpoint Threat Hunting, Malware Identification, Network Segmentation, VPN Controls, Botnet & Intrusion Analysis",
    "Software Development: React, TypeScript, Node.js, Vite, tRPC, Hono, REST APIs, Drizzle ORM, MySQL, TiDB Cloud, Vitest, HTML/Excel Dashboards",
    "Platforms & Leadership: Karmak DMS, Google Workspace, Microsoft Office, Dell PowerEdge Servers, Multi-Site IT Infrastructure, SLA Compliance, Vendor Coordination, Team Onboarding, Root Cause Analysis",
]

EDUCATION = "Associate of Arts, Communication Studies | American Public University System"


def _set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def build_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(CONTACT["name"].upper())
    _set_run_font(run, size=16, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(CONTACT["title"])
    _set_run_font(run, size=12)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = (
        f"{CONTACT['location']} | {CONTACT['phone']} | {CONTACT['email']} | "
        f"{CONTACT['github']}"
    )
    run = contact.add_run(contact_text)
    _set_run_font(run, size=10)

    def add_heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        _set_run_font(run, size=11, bold=True)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    def add_body(text: str) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)

    def add_bullet(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    add_heading("Professional Summary")
    add_body(SUMMARY)

    add_heading("Core Skills")
    for skill in SKILLS_ATS:
        add_bullet(skill)

    add_heading("Key Achievements")
    for item in ACHIEVEMENTS:
        add_bullet(item)

    add_heading("Professional Experience")
    for job in EXPERIENCE:
        header = doc.add_paragraph()
        run = header.add_run(job["title"])
        _set_run_font(run, bold=True)
        run = header.add_run(f" | {job['company']} | {job['location']} | {job['dates']}")
        _set_run_font(run)
        header.paragraph_format.space_after = Pt(2)
        for bullet in job["bullets"]:
            add_bullet(bullet)

    add_heading("Technical Projects")
    proj = doc.add_paragraph()
    run = proj.add_run(PROJECT["name"])
    _set_run_font(run, bold=True)
    run = proj.add_run(f" | {PROJECT['links']}")
    _set_run_font(run, size=10)
    for bullet in PROJECT["bullets"]:
        add_bullet(bullet)

    add_heading("Education")
    add_body(EDUCATION)

    doc.save(path)


def build_html(path: Path) -> None:
    achievements_html = "\n".join(f"<li>{a}</li>" for a in ACHIEVEMENTS)
    skills_html = "\n".join(f"<li>{s}</li>" for s in SKILLS_ATS)

    experience_html = ""
    for job in EXPERIENCE:
        bullets = "\n".join(f"<li>{b}</li>" for b in job["bullets"])
        experience_html += f"""
        <article class="job">
          <div class="job-header">
            <div>
              <h3>{job['title']}</h3>
              <p class="company">{job['company']}</p>
            </div>
            <div class="job-meta">
              <span>{job['location']}</span>
              <span>{job['dates']}</span>
            </div>
          </div>
          <ul>{bullets}</ul>
        </article>
        """

    project_bullets = "\n".join(f"<li>{b}</li>" for b in PROJECT["bullets"])

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{CONTACT['name']} - Resume</title>
  <style>
    @page {{
      size: letter;
      margin: 0.55in 0.65in;
    }}
    :root {{
      --ink: #1a2332;
      --muted: #4a5568;
      --accent: #1e5a8a;
      --accent-soft: #e8f1f8;
      --line: #d7dee8;
    }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: "Segoe UI", Calibri, Arial, sans-serif;
      color: var(--ink);
      line-height: 1.45;
      font-size: 10.5pt;
      background: #f4f7fb;
    }}
    .page {{
      max-width: 8.5in;
      margin: 0 auto;
      background: #fff;
      box-shadow: 0 8px 30px rgba(26, 35, 50, 0.08);
    }}
    header {{
      padding: 28px 34px 22px;
      border-bottom: 3px solid var(--accent);
      background: linear-gradient(180deg, #ffffff 0%, #f8fbff 100%);
    }}
    h1 {{
      font-size: 28pt;
      letter-spacing: 0.04em;
      font-weight: 700;
      color: var(--ink);
    }}
    .subtitle {{
      margin-top: 4px;
      font-size: 13pt;
      color: var(--accent);
      font-weight: 600;
    }}
    .contact {{
      margin-top: 12px;
      display: flex;
      flex-wrap: wrap;
      gap: 8px 18px;
      color: var(--muted);
      font-size: 9.5pt;
    }}
    .contact a {{
      color: var(--accent);
      text-decoration: none;
    }}
    main {{
      padding: 22px 34px 30px;
    }}
    section {{
      margin-bottom: 18px;
    }}
    h2 {{
      font-size: 11pt;
      text-transform: uppercase;
      letter-spacing: 0.08em;
      color: var(--accent);
      border-bottom: 1px solid var(--line);
      padding-bottom: 4px;
      margin-bottom: 10px;
    }}
    p {{ color: var(--ink); }}
    .summary {{ color: var(--muted); }}
    ul {{
      padding-left: 18px;
    }}
    li {{
      margin-bottom: 5px;
      color: var(--ink);
    }}
    .skills-grid {{
      display: grid;
      grid-template-columns: 1fr;
      gap: 6px;
      list-style: none;
      padding: 0;
    }}
    .skills-grid li {{
      background: var(--accent-soft);
      border-left: 3px solid var(--accent);
      padding: 7px 10px;
      border-radius: 4px;
      margin: 0;
      font-size: 9.8pt;
    }}
    .job {{
      margin-bottom: 14px;
      page-break-inside: avoid;
    }}
    .job-header {{
      display: flex;
      justify-content: space-between;
      gap: 16px;
      align-items: flex-start;
      margin-bottom: 6px;
    }}
    .job h3 {{
      font-size: 11pt;
      font-weight: 700;
      color: var(--ink);
    }}
    .company {{
      color: var(--muted);
      font-size: 10pt;
      margin-top: 2px;
    }}
    .job-meta {{
      text-align: right;
      color: var(--muted);
      font-size: 9.5pt;
      white-space: nowrap;
    }}
    .job-meta span {{
      display: block;
    }}
    .project-title {{
      font-weight: 700;
      margin-bottom: 4px;
    }}
    .project-links {{
      color: var(--muted);
      font-size: 9.5pt;
      margin-bottom: 6px;
    }}
    .project-links a {{
      color: var(--accent);
      text-decoration: none;
    }}
    .education {{
      font-weight: 600;
    }}
  </style>
</head>
<body>
  <div class="page">
    <header>
      <h1>{CONTACT['name'].upper()}</h1>
      <p class="subtitle">{CONTACT['title']}</p>
      <div class="contact">
        <span>{CONTACT['location']}</span>
        <span>{CONTACT['phone']}</span>
        <span><a href="mailto:{CONTACT['email']}">{CONTACT['email']}</a></span>
        <span><a href="{CONTACT['github']}">{CONTACT['github'].replace('https://', '')}</a></span>
        <span><a href="{CONTACT['portfolio']}">{CONTACT['portfolio'].replace('https://', '')}</a></span>
      </div>
    </header>
    <main>
      <section>
        <h2>Professional Summary</h2>
        <p class="summary">{SUMMARY}</p>
      </section>
      <section>
        <h2>Core Skills</h2>
        <ul class="skills-grid">{skills_html}</ul>
      </section>
      <section>
        <h2>Key Achievements</h2>
        <ul>{achievements_html}</ul>
      </section>
      <section>
        <h2>Professional Experience</h2>
        {experience_html}
      </section>
      <section>
        <h2>Technical Projects</h2>
        <p class="project-title">{PROJECT['name']}</p>
        <p class="project-links">
          <a href="https://github.com/RamenAnime/RamenAnime-Portfolio">github.com/RamenAnime/RamenAnime-Portfolio</a>
          &nbsp;|&nbsp;
          <a href="{CONTACT['portfolio']}">ramenanime.com</a>
        </p>
        <ul>{project_bullets}</ul>
      </section>
      <section>
        <h2>Education</h2>
        <p class="education">{EDUCATION}</p>
      </section>
    </main>
  </div>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def build_pdf(html_path: Path, pdf_path: Path) -> None:
    HTML(filename=str(html_path)).write_pdf(str(pdf_path))


def main() -> None:
    docx_path = OUT_DIR / "Jason_Jones_Resume_ATS.docx"
    html_path = OUT_DIR / "Jason_Jones_Resume_Visual.html"
    pdf_path = OUT_DIR / "Jason_Jones_Resume_Visual.pdf"

    build_docx(docx_path)
    build_html(html_path)
    build_pdf(html_path, pdf_path)

    print(f"Created: {docx_path}")
    print(f"Created: {html_path}")
    print(f"Created: {pdf_path}")


if __name__ == "__main__":
    main()
